"""
Export Service: generates PDF and DOCX from simplified notes and quizzes.
"""
import io
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER


def export_notes_pdf(filename: str, simplified_text: str, key_concepts: list,
                     complexity_level: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                 fontSize=16, spaceAfter=12, textColor=colors.HexColor("#1a3c6b"))
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"],
                                   fontSize=13, spaceAfter=8, textColor=colors.HexColor("#2c5282"))
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
                                fontSize=11, spaceAfter=6, leading=16)
    label_style = ParagraphStyle("Label", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.HexColor("#718096"))

    story = [
        Paragraph(f"Simplified Notes: {filename}", title_style),
        Paragraph(f"Complexity Level: {complexity_level.title()}", label_style),
        Spacer(1, 0.5*cm),
        Paragraph("Summary", heading_style),
    ]

    for para in simplified_text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), body_style))
            story.append(Spacer(1, 0.2*cm))

    if key_concepts:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("Key Concepts", heading_style))
        table_data = [["Term", "Definition"]]
        for kc in key_concepts:
            table_data.append([
                kc.get("term", ""),
                kc.get("definition", "—"),
            ])
        t = Table(table_data, colWidths=[5*cm, 12*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c5282")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#ebf4ff")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(t)

    doc.build(story)
    return buffer.getvalue()


def export_quiz_pdf(filename: str, questions: list) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                 fontSize=16, spaceAfter=12, textColor=colors.HexColor("#1a3c6b"))
    q_style = ParagraphStyle("Q", parent=styles["Normal"],
                             fontSize=11, spaceAfter=4, fontName="Helvetica-Bold")
    opt_style = ParagraphStyle("Opt", parent=styles["Normal"],
                               fontSize=11, spaceAfter=2, leftIndent=20)
    ans_style = ParagraphStyle("Ans", parent=styles["Normal"],
                               fontSize=10, spaceAfter=8, textColor=colors.HexColor("#276749"),
                               leftIndent=20)

    story = [Paragraph(f"Quiz: {filename}", title_style)]

    labels = ["A", "B", "C", "D"]
    for i, q in enumerate(questions, 1):
        story.append(Paragraph(f"Q{i}. {q['question']}", q_style))
        for j, opt in enumerate(q["options"]):
            story.append(Paragraph(f"{labels[j]}. {opt}", opt_style))
        correct_label = labels[q.get("correct_index", 0)]
        story.append(Paragraph(f"Answer: {correct_label}  — {q.get('explanation', '')}", ans_style))
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    return buffer.getvalue()


def export_notes_docx(filename: str, simplified_text: str, key_concepts: list,
                      complexity_level: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(f"Simplified Notes: {filename}", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x3c, 0x6b)

    label = doc.add_paragraph(f"Complexity Level: {complexity_level.title()}")
    label.runs[0].font.size = Pt(9)
    label.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_heading("Summary", level=2)
    for para in simplified_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    if key_concepts:
        doc.add_heading("Key Concepts", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Term"
        hdr[1].text = "Definition"
        for kc in key_concepts:
            row = table.add_row().cells
            row[0].text = kc.get("term", "")
            row[1].text = kc.get("definition", "—")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_quiz_docx(filename: str, questions: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(f"Quiz: {filename}", level=1)

    labels = ["A", "B", "C", "D"]
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{i}. {q['question']}")
        run.bold = True
        run.font.size = Pt(11)

        for j, opt in enumerate(q["options"]):
            doc.add_paragraph(f"  {labels[j]}. {opt}", style="List Bullet")

        correct_label = labels[q.get("correct_index", 0)]
        ans = doc.add_paragraph(f"Answer: {correct_label}  — {q.get('explanation', '')}")
        ans.runs[0].font.color.rgb = RGBColor(0x27, 0x67, 0x49)
        ans.runs[0].font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
